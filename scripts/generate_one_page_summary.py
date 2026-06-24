from pathlib import Path
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
metrics = json.loads((ROOT / "results" / "metrics.json").read_text(encoding="utf-8"))
out = Path(r"C:\Users\HP\Downloads\Adversarial_NIDS_One_Page_Summary_IITK_BCyber.docx")

doc = Document()
t = doc.add_paragraph("Adversarial Robustness in ML-based NIDS — One-Page Summary")
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].bold = True
t.runs[0].font.size = Pt(13)

rows = [
    ("Applicant", "Hitvika Pathak"),
    ("Project", "Adversarial Robustness in Network Intrusion Detection"),
    ("Dataset", metrics["dataset_file"]),
    ("Models", "Random Forest, XGBoost, MLP"),
    ("Attacks", "FGSM, PGD (white-box MLP + transfer to tree models)"),
    ("Defenses", "Adversarial training, feature squeezing"),
    (
        "Key result",
        f"RF clean {metrics['summary_table'][0]['clean_accuracy']}% → PGD transfer detection {metrics['summary_table'][0]['pgd_detection_rate']}%",
    ),
    ("Reproduce", "python scripts/run_experiments.py"),
    ("GitHub", "https://github.com/hitvika/adversarial-nids-cicids2017"),
    ("Local path", str(ROOT)),
]
for k, v in rows:
    p = doc.add_paragraph()
    p.add_run(f"{k}: ").bold = True
    p.add_run(v)

doc.save(out)
print(f"Saved: {out}")