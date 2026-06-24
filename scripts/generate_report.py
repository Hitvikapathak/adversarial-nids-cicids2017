"""Generate final IITK B.Cyber report with 17 required sections."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
OUT_DOCX = Path(r"C:\Users\HP\Downloads\Adversarial_Robustness_CIC_IDS2017_Project_Report_FINAL.docx")
OUT_PDF = Path(r"C:\Users\HP\Downloads\Adversarial_Robustness_CIC_IDS2017_Project_Report_FINAL.pdf")
REFLECTION = ROOT / "docs" / "personal_reflection.md"
GITHUB_REPO = "https://github.com/hitvika/adversarial-nids-cicids2017"


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph("")


def image(doc, path: Path, caption: str, width=5.3):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    m = json.loads((RESULTS / "metrics.json").read_text(encoding="utf-8"))
    s = m["summary_table"]
    profile = m["dataset_profile"]
    reflection = REFLECTION.read_text(encoding="utf-8") if REFLECTION.exists() else ""

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(m["project_title"])
    r.bold = True
    r.font.size = Pt(13)
    doc.add_paragraph("IIT Kanpur B.Cyber Programme — Project Report").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Hitvika Pathak | June 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, "Evidence of Work / Portfolio Summary")
    evidence = [
        ("GitHub repository", GITHUB_REPO),
        ("Local path", str(ROOT)),
        ("Reproduce", "python run.py"),
        ("Dataset", f"CIC-IDS2017 — {m['dataset_file']}"),
        ("Models", "Random Forest, XGBoost, MLP"),
        ("Attacks", "FGSM + PGD (white-box MLP, transfer to tree models)"),
        ("Defenses", "Adversarial training, feature squeezing, ensemble"),
        (
            "Key result",
            f"RF clean {s[0]['clean_accuracy']}% → PGD transfer detection {s[0]['pgd_detection_rate']}%",
        ),
        ("Demo video script", "docs/demo_video_script.md"),
    ]
    for k, v in evidence:
        doc.add_paragraph(f"{k}: {v}")

    heading(doc, "Abstract")
    doc.add_paragraph(
        "This project investigates whether machine-learning-based Network Intrusion Detection Systems remain "
        "reliable under adversarial manipulation. Using CIC-IDS2017 (Thursday Web Attacks subset, "
        f"{profile['rows']:,} raw flows), I trained Random Forest, XGBoost, and MLP classifiers and evaluated "
        "them with FGSM and PGD under a STRIDE threat model. High clean accuracy did not guarantee security: "
        f"Random Forest dropped from {s[0]['clean_accuracy']}% clean accuracy to "
        f"{s[0]['pgd_detection_rate']}% attack-detection rate under transferred PGD (ε=0.05). "
        "Adversarial training improved MLP PGD detection to "
        f"{m['defense_table'][1]['pgd_detection_rate']}%. The work shows why NIDS models must be tested under "
        "attacker-aware conditions before deployment."
    )

    heading(doc, "Problem Statement")
    doc.add_paragraph(
        "ML-NIDS classify network flows from statistical features. If an attacker can perturb those features "
        "within realistic bounds, malicious traffic may be labeled benign. This project measures that risk on "
        "CIC-IDS2017 and compares mitigation strategies."
    )

    heading(doc, "Why This Matters for Cybersecurity")
    for item in [
        "Evasion attacks can bypass automated alerts in SOC pipelines.",
        "Clean-test accuracy hides failure modes that appear only under adversarial input.",
        "Robustness checks should be integrated into secure SDLC for detection systems.",
        "Layered defense (ML + constraints + logging + analyst review) is required in production.",
        "For national infrastructure, robust detection is a security control, not a leaderboard metric.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    heading(doc, "Dataset and Preprocessing")
    doc.add_paragraph(
        f"Source: CIC-IDS2017 ({m['dataset_file']}). Raw records: {profile['rows']:,}; features used: "
        f"{profile['feature_count']}; missing values handled: {profile['missing_values']:,}; duplicates removed in "
        f"analysis: {profile['duplicate_rows']:,}. Labels collapsed to binary Benign vs Attack. "
        f"Class distribution (raw): {profile['label_distribution']}. "
        f"Preprocessing: drop identifiers, numeric coercion, top-30 variance features, StandardScaler, "
        f"balanced sampling ({m['samples_after_balancing']} flows), split {m['splits']}."
    )

    heading(doc, "Threat Model using STRIDE")
    table(
        doc,
        ["STRIDE", "NIDS Example", "Feature-Level Attack", "Defense"],
        [
            ["Tampering", "Alter timing behavior", "Flow Duration / IAT perturbation", "Constraint validation"],
            ["Spoofing", "Mimic benign traffic", "Reduce packet-rate features", "Behavioral baselines"],
            ["Repudiation", "Hide malicious footprint", "Suppress volume features", "Immutable flow logs"],
            ["Information Disclosure", "Model probing", "Boundary search via queries", "Access control"],
            ["Denial of Service", "Flood crafted flows", "High-volume adversarial samples", "Rate limits + robust model"],
            ["Elevation of Privilege", "Bypass policy action", "Force benign classification", "Defense-in-depth rules"],
        ],
    )
    heading(doc, "Realistic Perturbation Constraints", 2)
    for c in [
        "Packet counts and durations remain non-negative.",
        "Protocol identifiers are not arbitrarily changed.",
        "Derived statistics stay internally consistent.",
        "Perturbations preserve attack intent (evasion, not benign conversion).",
        "Categorical network fields are not treated as free continuous pixels.",
        "All perturbations clipped to L∞ bounds after scaling.",
    ]:
        doc.add_paragraph(c, style="List Bullet")

    heading(doc, "Models Implemented")
    doc.add_paragraph(
        "Random Forest (200 trees, depth 20, balanced class weights) for strong tabular baseline and interpretability. "
        "XGBoost (200 estimators) for boosted nonlinear boundaries. MLP (64-32 ReLU, Adam) as differentiable "
        "surrogate for gradient attacks."
    )

    heading(doc, "Adversarial Attack Design")
    doc.add_paragraph(
        "FGSM and PGD require gradients, so white-box attacks were executed on MLP using finite-difference gradients "
        "over attack-class loss. Random Forest and XGBoost were evaluated with transfer attacks from MLP-generated "
        "adversarial examples."
    )
    table(
        doc,
        ["Attack", "Epsilon", "Steps", "Step Size", "Targeted?", "Model"],
        [
            ["FGSM", "0.01, 0.05, 0.1", "1", "—", "No", "MLP (white-box)"],
            ["PGD", "0.01, 0.05, 0.1", str(m["pgd_steps"]), str(m["pgd_step_size"]), "No", "MLP (white-box)"],
            ["Transfer", "0.05 primary", "—", "—", "No", "RF, XGBoost"],
        ],
    )

    heading(doc, "Defense Mechanisms")
    doc.add_paragraph(
        "1) Adversarial training: MLP retrained on clean + PGD examples (ε=0.05). "
        "2) Feature squeezing: 4-bit precision reduction before inference. "
        "3) Ensemble: majority vote of Random Forest and robust MLP."
    )

    heading(doc, "Experimental Setup and Reproducibility")
    doc.add_paragraph(
        f"Seed: {m['random_seed']}. Environment: Python 3.12, Windows 10, CPU execution. "
        f"Attack evaluation set: {m['attack_samples_evaluated']} attack flows. "
        "Command: python run.py. Artifacts: results/metrics.json, results/*.png, models/all_models.joblib."
    )

    heading(doc, "Results and Graphs")
    doc.add_paragraph("PGD/FGSM columns report attack-detection rate on adversarial attack flows.")
    table(
        doc,
        ["Model", "Clean Acc", "Clean Recall", "F1", "FGSM Det.", "PGD Det.", "Post-Defense PGD"],
        [
            [
                r["model"],
                f"{r['clean_accuracy']}%",
                f"{r['clean_recall']}%",
                f"{r['clean_f1']}%",
                f"{r['fgsm_detection_rate']}%",
                f"{r['pgd_detection_rate']}%",
                "—" if r["after_defense_pgd_rate"] is None else f"{r['after_defense_pgd_rate']}%",
            ]
            for r in s
        ],
    )
    table(
        doc,
        ["Defense", "Clean Acc", "PGD Det.", "Compute Cost"],
        [[d["defense"], f"{d['clean_accuracy']}%", f"{d['pgd_detection_rate']}%", d["compute_cost"]] for d in m["defense_table"]],
    )
    image(doc, RESULTS / "accuracy_vs_epsilon.png", "Figure 1: Attack-detection rate vs epsilon.")
    image(doc, RESULTS / "accuracy_drop_comparison.png", "Figure 2: Clean vs PGD performance by model.")
    image(doc, RESULTS / "cm_random_forest_pgd.png", "Figure 3: Random Forest under transferred PGD.")
    doc.add_paragraph(
        f"At ε=0.1, MLP detection fell to "
        f"{[x for x in m['epsilon_curve'] if x['model']=='MLP' and x['epsilon']==0.1][0]['attack_detection_rate']}%, "
        "showing sensitivity to larger perturbation budgets."
    )

    heading(doc, "Failure Analysis")
    for f in [
        f"Transferred PGD was most damaging to tree models ({s[0]['pgd_detection_rate']}% RF detection).",
        f"MLP white-box PGD at ε=0.05 retained {s[2]['pgd_detection_rate']}% detection but degraded at higher epsilon.",
        "Feature squeezing did not improve PGD detection in this run.",
        "Class imbalance mitigation helped clean recall, but attack-only adversarial recall remained the critical metric.",
        f"Adversarial training improved MLP PGD detection to {m['defense_table'][1]['pgd_detection_rate']}% with minor clean-accuracy trade-off.",
    ]:
        doc.add_paragraph(f, style="List Bullet")

    heading(doc, "Deployment Considerations")
    doc.add_paragraph(
        "Deploy with feature constraint checks, dual-model ensemble scoring, centralized logging, periodic red-team "
        "adversarial tests, and human analyst escalation for low-confidence alerts."
    )

    heading(doc, "Limitations")
    doc.add_paragraph(
        "Single-day CIC-IDS2017 file, sampled flows for compute limits, feature-level attacks (not packet-level), "
        "and finite-difference gradients instead of native autograd."
    )

    heading(doc, "Future Work")
    heading(doc, "Extension: From Model Robustness to Operational Security", 2)
    for item in [
        "FastAPI dashboard for live flow classification.",
        "API security middleware for inference endpoint protection.",
        "Cloud deployment demo with monitored endpoints.",
        "Threat-analysis pipeline from traffic logs.",
        "CTF-style adversarial evasion challenge.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    heading(doc, "References")
    refs = [
        "Goodfellow, I. J., Shlens, J., & Szegedy, C. (2015). Explaining and Harnessing Adversarial Examples.",
        "Madry, A., et al. (2018). Towards Deep Learning Models Resistant to Adversarial Attacks.",
        "Carlini, N., & Wagner, D. (2017). Towards Evaluating the Robustness of Neural Networks.",
        "Xu, W., Evans, D., & Qi, Y. (2018). Feature Squeezing.",
        "Sharafaldin, I., Lashkari, A. H., & Ghorbani, A. A. (2018). CIC-IDS2017 dataset paper.",
        "Nicolae, I., et al. (2018). Adversarial Robustness Toolbox.",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f"[{i}] {ref}")

    heading(doc, "Appendix: GitHub, Screenshots, Code Snippets")
    doc.add_paragraph(f"GitHub: {GITHUB_REPO}")
    doc.add_paragraph(f"Repository root: {ROOT}")
    doc.add_paragraph("Reproduce: python run.py")
    doc.add_paragraph("Screenshot artifact: results/screenshots/terminal_output.png")
    image(doc, RESULTS / "screenshots" / "terminal_output.png", "Appendix Figure: Pipeline terminal output.", 6.0)
    doc.add_paragraph(
        "Code snippet (attack call): pgd_adv = pgd_attack(mlp, attack_x, attack_y, epsilon=0.05)"
    )
    if reflection:
        heading(doc, "Personal Reflection", 2)
        doc.add_paragraph(reflection.replace("# Personal Reflection — Hitvika Pathak\n\n", ""))

    doc.save(OUT_DOCX)
    print(f"Saved: {OUT_DOCX}")

    # Best-effort PDF export
    try:
        import subprocess

        soffice = Path(r"C:\Users\HP\.grok\skills\docx\scripts\office\soffice.py")
        if soffice.exists():
            subprocess.run(
                ["python", str(soffice), "--headless", "--convert-to", "pdf", str(OUT_DOCX)],
                check=True,
                cwd=str(OUT_DOCX.parent),
            )
            print(f"Saved: {OUT_PDF}")
    except Exception as exc:
        print(f"PDF export skipped: {exc}")


if __name__ == "__main__":
    main()